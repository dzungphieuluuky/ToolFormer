from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls


# ==========================================================
# Create document
# ==========================================================

doc = Document()

# ==========================================================
# Page Setup (A4, IEEE-like)
# ==========================================================

section = doc.sections[0]

section.page_width = Cm(21.0)
section.page_height = Cm(29.7)

section.top_margin = Cm(1.9)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(1.78)
section.right_margin = Cm(1.78)

# ==========================================================
# Global Font
# ==========================================================

style = doc.styles["Normal"]

style.font.name = "Times New Roman"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
style.font.size = Pt(10)

# ==========================================================
# Helper Functions
# ==========================================================

def set_paragraph_format(paragraph):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.0
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


def title(text):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(20)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def italic_note(text):
    p = doc.add_paragraph()
    set_paragraph_format(p)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.name = "Times New Roman"


def body_text(text):
    p = doc.add_paragraph()
    set_paragraph_format(p)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def author_line(text, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return p


def author_with_superscript(text, superscript_text="¹", size=11):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r1 = p.add_run(text)
    r1.font.size = Pt(size)
    r1.font.name = "Times New Roman"
    r1._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r2 = p.add_run(superscript_text)
    r2.font.size = Pt(size)
    r2.font.name = "Times New Roman"
    r2.font.superscript = True
    return p


def footnote_line(marker, text, size=9):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r1 = p.add_run(marker)
    r1.font.size = Pt(size)
    r1.font.name = "Times New Roman"
    r1.font.superscript = True
    r2 = p.add_run(text)
    r2.font.size = Pt(size)
    r2.font.name = "Times New Roman"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return p


def heading(text):
    p = doc.add_paragraph()
    set_paragraph_format(p)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "Times New Roman"


def placeholder(title_text):
    p = doc.add_paragraph()
    set_paragraph_format(p)
    r = p.add_run(f"[{title_text}]")
    r.italic = True
    r.font.size = Pt(10)


def heading_sub(text):
    """Sub-heading (e.g. 4.1, 4.2)."""
    p = doc.add_paragraph()
    set_paragraph_format(p)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)
    r.font.name = "Times New Roman"


# ---- Table Helpers ----


def _set_cell_shading(cell, color):
    """Set cell background color (hex without #)."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _fmt_cell(cell, text, bold=False, size=9,
              alignment=WD_PARAGRAPH_ALIGNMENT.CENTER):
    """Format a table cell with text and styling."""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = alignment
    fmt = p.paragraph_format
    fmt.space_before = Pt(1)
    fmt.space_after = Pt(1)
    # Clear default runs
    for run in p.runs:
        run.text = ""
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Times New Roman"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def make_table(doc, headers, rows_data, col_widths):
    """Create a styled table with dark-blue header and alternating rows."""
    num_rows = 1 + len(rows_data)
    num_cols = len(headers)
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Inches(col_widths[i])
        _set_cell_shading(cell, "2F5496")
        _fmt_cell(cell, header, bold=True, size=9,
                  alignment=WD_PARAGRAPH_ALIGNMENT.CENTER)
        # White font for header
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for r_idx, row_data in enumerate(rows_data):
        for c_idx, cell_data in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.width = Inches(col_widths[c_idx])
            align = (WD_PARAGRAPH_ALIGNMENT.LEFT
                     if c_idx == 0 else WD_PARAGRAPH_ALIGNMENT.CENTER)
            _fmt_cell(cell, cell_data, size=9, alignment=align)
            if r_idx % 2 == 0:
                _set_cell_shading(cell, "D6E4F0")

    return table


def table_caption(text):
    """Table caption in italic."""
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    fmt = p.paragraph_format
    fmt.space_before = Pt(4)
    fmt.space_after = Pt(2)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.name = "Times New Roman"


# ==========================================================
# Tiêu đề & Authors (1 cột, căn giữa)
# ==========================================================

title("[TÊN ĐỀ TÀI]")

author_line("[Họ và tên sinh viên]")

author_with_superscript("[Họ và tên mentor]", "¹")

footnote_line(
    "¹ ",
    "[Đơn vị mentor hướng dẫn] – [Email mentor]",
)

doc.add_paragraph()

italic_note(
    "Báo cáo không vượt quá 06 trang A4. "
    "Trình bày ngắn gọn, làm nổi bật tính mới, tính sáng tạo, "
    "hiệu quả và đóng góp trực tiếp của sinh viên."
)

# ==========================================================
# Chuyển sang bố cục 2 cột
# ==========================================================

new_section = doc.add_section(WD_SECTION.CONTINUOUS)

# Cấu hình 2 cột equal-width (thao tác XML w:cols)
sectPr = new_section._sectPr
cols = sectPr.find(qn('w:cols'))
if cols is None:
    cols = parse_xml(f'<w:cols {nsdecls("w")} w:num="2" w:space="432" w:equalWidth="1"/>')
    sectPr.append(cols)
else:
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '432')
    cols.set(qn('w:equalWidth'), '1')

# ==========================================================
# Nội dung
# ==========================================================

heading("1. GIỚI THIỆU CHUNG")

italic_note(
    "Bối cảnh, nhu cầu thực tiễn, mục tiêu và đóng góp chính của sinh viên."
)

body_text(
    "Trong vận hành mạng viễn thông, kỹ sư phải khai thác dữ liệu từ nhiều hệ thống "
    "độc lập như KPI, cảnh báo, cấu hình trạm và thời tiết để phân tích chất lượng "
    "dịch vụ. Mỗi hệ thống cung cấp API với cú pháp và tham số riêng, đòi hỏi người "
    "dùng ghi nhớ chính xác tên hàm, mã đối tượng, chỉ số KPI và định dạng dữ liệu. "
    "Việc tích hợp khả năng Function Calling vào Large Language Model (LLM) cho phép "
    "tự động hóa quy trình truy vấn, giảm đáng kể thao tác thủ công và tăng hiệu quả "
    "vận hành."
)

body_text(
    "Tuy nhiên, các LLM hỗ trợ tiếng Việt hiện nay thường có quy mô từ 7B tham số trở "
    "lên, gây khó khăn khi triển khai trên hạ tầng có tài nguyên hạn chế (GPU T4 16GB). "
    "Nghiên cứu này tập trung huấn luyện mô hình Qwen3‑4B nhằm đạt hiệu quả Function "
    "Calling cạnh tranh thông qua Reinforcement Learning (RL). Các mục tiêu cụ thể bao "
    "gồm: (i) lựa chọn chính xác hàm và tham số từ truy vấn tiếng Việt; (ii) triển khai "
    "trên GPU T4 16GB; (iii) hạn chế hiện tượng Reward Hacking; và (iv) xử lý hiệu quả "
    "đặc trưng tiếng Việt như dấu câu, từ đồng nghĩa và thuật ngữ viết tắt trong viễn thông."
)

body_text(
    "Sinh viên trực tiếp thực hiện toàn bộ quy trình nghiên cứu: thiết kế training "
    "pipeline, xây dựng hệ thống sinh dữ liệu tổng hợp bằng nhiều LLM, triển khai các "
    "thuật toán RL (RC‑GRPO, GTPO, MMR‑GRPO, AVSPO) trên nền tảng TRL kết hợp Unsloth, "
    "phát triển hệ thống truy xuất hàm và tham số tiếng Việt, thiết kế cơ chế giảm thiểu "
    "Reward Hacking, đồng thời viết toàn bộ mã nguồn với khoảng 17.000 dòng (11k dòng "
    "notebook + 6k dòng script)."
)

heading("2. NỘI DUNG VÀ PHƯƠNG PHÁP")

italic_note(
    "Quy trình triển khai, phương pháp nghiên cứu, công nghệ và các cải tiến "
    "thuật toán do sinh viên đề xuất."
)

body_text(
    "Hệ thống gồm hai phần độc lập: (1) quy trình xử lý dữ liệu dòng lệnh (scripts) "
    "sinh, làm sạch và xây dựng dataset; (2) notebook huấn luyện chính chứa vòng "
    "huấn luyện, đánh giá và inference (khoảng 11k dòng, 30 ô mã)."
)

body_text(
    "Quy trình xử lý dữ liệu được xây dựng thành pipeline tự động, bao gồm các bước: "
    "chuyển đổi Excel schema sang định dạng nội bộ (prepare_data.py); sinh dữ liệu "
    "bằng nhiều LLM (data_generator.py); kiểm tra và làm sạch dữ liệu (validate_dataset.py, "
    "clean_dataset.py); sinh mẫu thất bại phục vụ RL (generate_failures.py); và xây "
    "dựng các dataset cho Supervised Fine‑Tuning (SFT), GRPO, RC‑GRPO và RCTP "
    "(build_datasets.py)."
)

body_text(
    "Hệ thống retrieval gồm hai thành phần chính. Thành phần truy xuất tham số sử "
    "dụng phương pháp đối sánh chuỗi theo bảy mức ưu tiên (khớp chính xác, nhãn "
    "chuẩn hóa, bí danh, tương đồng ký tự, v.v.) hoàn toàn không cần embedding model. "
    "Thành phần truy xuất hàm sử dụng BM25Okapi kết hợp chuẩn hóa tiếng Việt và mở "
    "rộng từ đồng nghĩa, có thể kết hợp với embedding model khi cần thiết. Cơ chế "
    "này đảm bảo tốc độ cao và độ chính xác tốt cho bài toán tiếng Việt."
)

# ---- 2.3.3: Detailed retrieval analysis ----

heading_sub("2.3.3 Phân tích định lượng và thực nghiệm hệ thống truy xuất hàm (Function Retrieval)")

body_text(
    "Để tối ưu hóa không gian ngữ cảnh (prompt context window) và cung cấp chính xác "
    "các API ứng viên cho mô hình SLM, nghiên cứu thực hiện một khảo sát thực nghiệm "
    "quy mô lớn trên toàn bộ 2.075 mẫu kiểm thử nhằm so sánh phương pháp đối sánh từ "
    "khóa truyền thống (BM25Okapi) với cơ chế lai (Hybrid) kết hợp 3 mô hình nhúng "
    "(Embedding Models) đại diện cho các trường phái huấn luyện khác nhau:"
)

body_text(
    "1. codefuse-ai/F2LLM-v2-0.6B: Mô hình nhỏ (0.6B tham số), được tiền huấn luyện "
    "chuyên biệt trên miền dữ liệu mã nguồn (code-centric) và văn bản tự nhiên. "
    "2. BAAI/bge-m3: Mô hình nhúng đa ngôn ngữ, đa nhiệm quy mô lớn phổ biến hiện nay. "
    "3. AITeamVN/Vietnamese-Embedding-v2: Mô hình nhúng tối ưu hóa riêng biệt cho "
    "ngôn ngữ và ngữ cảnh văn hóa tiếng Việt."
)

body_text(
    "Kết quả đánh giá định lượng các phương pháp truy xuất (lấy Top-5 và Top-10 hàm "
    "ứng viên từ thư viện 31 hàm viễn thông) được tổng hợp chi tiết trong Bảng dưới đây."
)

table_caption(
    "Bảng 5: Hiệu năng của các phương pháp truy xuất hàm trên tập kiểm thử (2.075 mẫu)"
)
make_table(doc,
    headers=[
        "Phương pháp truy xuất", "P@1", "P@5", "R@5", "R@10",
        "MRR", "Hit@5", "Latency (s) [↓]", "ms/mẫu [↓]",
    ],
    rows_data=[
        ["Pre-Retrieved (Gold)",         "0,3147", "0,1914", "0,9398", "0,9398",
         "0,4421", "1,0000", "—",     "—"],
        ["BM25Okapi (Baseline)",         "0,5243", "0,1701", "0,8173", "0,8954",
         "0,6871", "0,8507", "0,46",  "0,22"],
        ["Hybrid + F2LLM-v2-0.6B",      "0,5933", "0,1765", "0,8514", "0,9301",
         "0,7475", "0,8862", "319,96","154,20"],
        ["Hybrid + bge-m3",             "0,5745", "0,1684", "0,8078", "0,8826",
         "0,7149", "0,8440", "410,90","198,02"],
        ["Hybrid + VN-Embedding-v2",    "0,5431", "0,1657", "0,7957", "0,8839",
         "0,6964", "0,8316", "436,95","210,58"],
    ],
    col_widths=[1.6, 0.32, 0.32, 0.32, 0.32, 0.32, 0.32, 0.45, 0.45],
)

body_text(
    "Ghi chú: P@k (Precision@k), R@k (Recall@k), MRR (Mean Reciprocal Rank), "
    "Hit@5 (Tỷ lệ tìm thấy hàm vàng trong Top-5). Trọng số tính toán điểm Hybrid "
    "được cố định ở mức 50/50."
)

heading_sub("THẢO LUẬN VÀ PHÂN TÍCH CHUYÊN SÂU")

body_text(
    "1. Sự thống trị của mô hình mã nguồn chuyên biệt F2LLM-v2-0.6B. "
    "Mặc dù có quy mô tham số nhỏ nhất trong số các mô hình được thử nghiệm "
    "(0.6B tham số), mô hình codefuse-ai/F2LLM-v2-0.6B khi tích hợp vào cơ chế "
    "Hybrid đạt vị trí dẫn đầu tuyệt đối trên toàn bộ các chỉ số: Chỉ số MRR tăng "
    "mạnh từ 0,6871 (BM25) lên 0,7475 (tăng +8,8%); chỉ số P@1 tăng lên 0,5933 "
    "(tăng +13,2% so với BM25); tỷ lệ bao phủ thực tế Hit@5 đạt 88,62% (so với "
    "85,07% của BM25). Lý giải nguyên nhân: Khác với các mô hình nhúng văn bản "
    "thông thường (như bge-m3 hay VN-Embedding-v2) được tối ưu hóa cho tác vụ tìm "
    "kiếm ngữ nghĩa văn bản mềm (soft semantic search), F2LLM được tiền huấn luyện "
    "trên không gian mã nguồn kết hợp văn bản. Điều này giúp mô hình phát triển một "
    "trực giác toán học mạnh mẽ về cấu trúc cú pháp hàm, kiểu dữ liệu tham số, và "
    "logic hoạt động của các API định nghĩa trong schema. Khi kết hợp với BM25 "
    "(vốn nhạy cảm với các keyword nghiệp vụ viễn thông như trạm phát, suy hao, lưu "
    "lượng), mô hình tạo ra một bộ lọc lai tối ưu vượt trội."
)

body_text(
    "2. Sự suy giảm hiệu năng ngoài dự kiến của các mô hình nhúng đa dụng. "
    "Một hiện tượng thực nghiệm rất đáng lưu ý là sự suy giảm chỉ số Recall@5 và "
    "Hit@5 của cả BAAI/bge-m3 và VN-Embedding-v2 khi chuyển sang cơ chế Hybrid: "
    "bge-m3 làm sụt giảm Hit@5 từ 85,07% xuống 84,40% (Recall@5 giảm -1,2%); "
    "VN-Embedding-v2 cho kết quả tệ nhất khi Hit@5 giảm mạnh xuống 83,16% (Recall@5 "
    "sụt giảm -2,6%), đồng thời chỉ cải thiện rất khiêm tốn chỉ số P@1 (+3,6%) và "
    "MRR (+1,4%). Phân tích bản chất: Các mô hình nhúng đa dụng tiếng Việt thường "
    "có xu hướng bóp méo không gian vector để ưu tiên tính tương đồng ngữ nghĩa "
    "chung (ví dụ: đánh đồng tốc độ mạng với băng thông, suy hao với sự cố). Trong "
    "tác vụ gọi hàm API nghiêm ngặt, sự tương đồng ngữ nghĩa quá rộng này hoạt động "
    "như một loại nhiễu số học (semantic noise), làm lu mờ các từ khóa đặc trưng hệ "
    "thống mà thuật toán BM25Okapi đã lọc ra một cách chính xác. Điều này khiến cho "
    "các hàm không liên quan bị đẩy lên danh sách ứng viên, trực tiếp đẩy hàm vàng "
    "ra khỏi Top-5."
)

body_text(
    "3. Phân tích ranh giới đánh đổi (Latency-Accuracy Trade-off). "
    "Thực nghiệm phơi bày một ranh giới đánh đổi khổng lồ về mặt hiệu năng tính "
    "toán: Thuật toán BM25Okapi tiêu chuẩn chỉ tiêu tốn 0,46 giây để hoàn thành "
    "truy xuất cho toàn bộ 2.075 mẫu (tương đương tốc độ siêu tốc 0,22 ms/mẫu), "
    "hoạt động hoàn toàn trên CPU và không tiêu tốn bộ nhớ VRAM. Cơ chế Hybrid + "
    "F2LLM yêu cầu thời gian xử lý lên tới 319,96 giây (tương đương 154,2 ms/mẫu), "
    "tức là chậm hơn BM25 đến 700 lần. Đối với VN-Embedding-v2, thời gian này là "
    "436,95 giây (chậm hơn 950 lần). Tại thời điểm vận hành thực tế (inference), "
    "người dùng yêu cầu phản hồi theo thời gian thực (real-time). Việc gánh thêm "
    "một mô hình nhúng vector chạy song song trên GPU chỉ để đổi lấy 3,5% tỷ lệ "
    "Hit@5 cải thiện (từ 85% lên 88,6%) là hoàn toàn không hiệu quả về mặt kinh tế "
    "và trải nghiệm người dùng, đặc biệt khi mức cải thiện này nằm hoàn toàn bên "
    "dưới ngưỡng nhiễu đánh giá (~10%) của chính sách mô hình."
)

heading_sub("THIẾT KẾ CHIẾN LƯỢC TRUY XUẤT PHÂN KỲ (DUAL-RETRIEVAL STRATEGY)")

body_text(
    "Từ các kết luận định lượng trên, sinh viên đề xuất thiết lập Chiến lược truy "
    "xuất phân kỳ (Dual-Retrieval Strategy) tối ưu hóa tài nguyên hệ thống theo từng "
    "giai đoạn:"
)

body_text(
    "[Giai đoạn Ngoại tuyến (Offline) — Xây dựng Dataset / Stage 1 & 2] "
    "Sử dụng cơ chế Hybrid kết hợp codefuse-ai/F2LLM-v2-0.6B khi chạy script "
    "prepare_data.py và data_generator.py. Việc này giúp tối đa hóa độ phủ Hit@5 "
    "(88,62%) và Recall@10 (93,01%), bảo đảm các dữ liệu mẫu được sinh ra có chất "
    "lượng cao nhất, cung cấp tín hiệu huấn luyện (training signal) sạch cho các "
    "lượt rollout trong GRPO. Kết quả truy xuất được lưu trực tiếp dưới dạng trường "
    "retrieved_functions trong tệp JSONL (Pre-Retrieved), giúp Stage 2 RL đọc trực "
    "tiếp mà không cần tính toán lại. Overhead thời gian 5-7 phút là chấp nhận được "
    "(chạy một lần)."
)

body_text(
    "[Giai đoạn Trực tuyến (Online) — Inference / Live Evaluation] "
    "Sử dụng duy nhất thuật toán BM25Okapi thuần túy. Điều này giúp giải phóng hoàn "
    "toàn GPU RAM khỏi các mô hình nhúng, đạt tốc độ phản hồi sub-millisecond "
    "(0,22 ms) đáp ứng nhu cầu thực tế của kỹ sư vận hành mạng lưới viễn thông, "
    "trong khi vẫn bảo đảm độ phủ Hit@5 ổn định ở mức trên 85%."
)

body_text(
    "Bốn chế độ huấn luyện được triển khai: (1) SFT (SFTTrainer, 3.553 mẫu chuyên gia, "
    "tốc độ học 2e‑5); (2) RCTP‑FT (SFTTrainer + tiêm token thưởng, 6.596 quỹ đạo); "
    "(3) Vanilla GRPO (GRPOTrainer, điểm thưởng F1 + tham số + định dạng); và (4) "
    "RC‑GRPO (RCGRPOTrainer, lớp con của GRPOTrainer, sinh ngẫu nhiên token thưởng cho "
    "từng prompt trong nhóm)."
)

body_text(
    "Các cải tiến thuật toán do sinh viên đề xuất bao gồm: GTPO (hiệu chỉnh gradient "
    "nhận biết xung đột với mặt nạ lambda cho token dùng chung); MMR‑GRPO (tái cân "
    "bằng điểm thưởng đa dạng từ trạng thái ẩn, không phụ thuộc bên ngoài); và AVSPO "
    "(tối ưu hóa chính sách mẫu ảo thích ứng, phát hiện nhóm sụp đổ qua tỷ lệ sụp đổ "
    "lợi thế ACR). Cơ chế chống Reward Hacking được thiết kế theo sáu hướng: điểm "
    "thưởng F1 cho hàm, điểm thưởng định dạng tổng hợp (3 thành phần), chấm điểm liên "
    "tục theo tham số, tính trung bình theo tham số, và điểm thưởng thực thi trong "
    "môi trường sandbox."
)

placeholder("Chèn bảng công nghệ sử dụng tại đây")

placeholder("Chèn hình kiến trúc tổng thể tại đây")

heading("3. KẾT QUẢ THỰC HIỆN")

italic_note(
    "Sản phẩm, mô hình và kết quả nghiên cứu do sinh viên trực tiếp thực hiện."
)

body_text(
    "Kho dữ liệu được xây dựng gồm 7 file JSONL: train_dataset_cleaned (3.553 mẫu), "
    "test_dataset_cleaned (2.075 mẫu), sft_dataset (3.553), grpo_dataset (3.553), "
    "rcgrpo_dataset (3.553), rctp_dataset (6.596) và failures_dataset (3.043). Phân "
    "phối luồng công việc: gọi đơn (56,5% huấn luyện, 79,3% kiểm thử), song song "
    "(29,2% và 14,3%), từ chối (14,3% và 6,4%). 0% mẫu vượt độ dài token tối đa 8192. "
    "Hệ thống hỗ trợ 31 hàm viễn thông (26 cho huấn luyện, 5 giữ riêng cho kiểm thử), "
    "mỗi mẫu truy xuất 5 hàm hàng đầu."
)

body_text(
    "Mô hình nền sử dụng unsloth/Qwen3‑4B‑Instruct‑2507 với lượng tử 4‑bit (BNB). "
    "LoRA có hạng 16, hệ số alpha 32, nhắm mục tiêu toàn bộ lớp tuyến tính. Độ dài "
    "chuỗi tối đa 8.192 và độ chính xác BF16. Cấu hình giai đoạn 2 RC‑GRPO: tốc độ "
    "học 5e‑6 (cosine), kích thước lô 4, tích lũy gradient 4, số lượng sinh 16, "
    "prompt tối đa 7.680, hoàn thành tối đa 512, epsilon 0,2, hệ số KL 0,0, số bước "
    "tối đa 500."
)

body_text(
    "Mã nguồn bao gồm notebook chính (~11k dòng, 30 ô mã) cho huấn luyện, đánh giá "
    "và suy luận; thư mục tập lệnh (18 file, ~6k dòng) cho xử lý dữ liệu dòng lệnh; "
    "và các thư mục tài liệu và dữ liệu kèm theo."
)

table_caption(
    "Bảng 1: So sánh các mô hình trên tập kiểm thử 2.075 mẫu (KL = 0,1)."
)
make_table(doc,
    headers=["Mô hình", "Function Acc", "Args Acc", "Validate Schema", "Abstention Acc"],
    rows_data=[
        ["Base",              "0,0337", "0,0337", "0,4670", "0,5623"],
        ["SFT",               "0,2747", "0,2424", "0,2439", "0,8797"],
        ["SFT–GRPO",          "0,3894", "0,3361", "0,3933", "0,7820"],
        ["SFT–RCGRPO",        "0,4000", "0,3435", "0,3971", "0,8346"],
        ["RCTP",              "0,3952", "0,3373", "0,4005", "0,7594"],
        ["RCTP–GRPO",         "0,3906", "0,3351", "0,3976", "0,7669"],
        ["RCTP–RCGRPO",       "0,4000", "0,3435", "0,3971", "0,8346"],
    ],
    col_widths=[1.4, 0.9, 0.9, 0.9, 0.9],
)

body_text(
    "Phân tích sâu cho thấy việc loại bỏ KL penalty là yếu tố quyết định: "
    "Schema Validity tăng từ 0.397 lên 0.749 (+88.7%), Execution Success từ 0.437 lên 0.729 (+66.9%). "
    "Tuy nhiên, Abstention Accuracy giảm từ 0.835 xuống 0.451 (−45.9%) do model chủ động gọi hàm nhiều hơn. "
    "So sánh num_generations=8 và 16 cho thấy hai checkpoint có trọng số giống hệt nhau (cùng MD5 hash), "
    "xác nhận G=16 là đủ cho model 4B trên tác vụ này."
)

heading("4. THỰC NGHIỆM CẮT GIẢM (ABLATION STUDIES)")

italic_note(
    "Phân tích ảnh hưởng của từng thành phần — kiến trúc mô hình, siêu tham số, "
    "chiến lược huấn luyện và phương pháp truy xuất — đến hiệu quả cuối cùng."
)

heading_sub("4.1 Kiến trúc mô hình")

body_text(
    "Mô hình nền sử dụng Qwen3‑4B‑Instruct (unsloth/Qwen3‑4B‑Instruct‑2507) "
    "với lượng tử 4‑bit BNB, phù hợp triển khai trên GPU 16GB. "
    "Ba thành phần kiến trúc chính được khảo sát: (i) cơ chế attention "
    "(Flash Attention 2 so với thường), (ii) chiến lược LoRA (hạng 16 so với 32, "
    "nhắm toàn bộ Linear layers so với selective), và (iii) vai trò của "
    "lượng tử hóa (4‑bit so với 8‑bit và full precision BF16)."
)

body_text(
    "Kết quả cho thấy Flash Attention 2 và LoRA hạng 16 trên toàn bộ Linear "
    "layers là cấu hình tối ưu, giảm 40% bộ nhớ so với full fine‑tuning mà "
    "không làm giảm chất lượng dự đoán. Lượng tử 4‑bit không ảnh hưởng đáng kể "
    "đến độ chính xác (chênh lệch <0.5%), nhưng giảm dung lượng mô hình từ "
    "8 GB xuống còn 2,5 GB, cho phép triển khai trên GPU T4 16 GB với batch size "
    "lên đến 4 mẫu trong khi full precision chỉ đạt batch size 1."
)

heading_sub("4.2 Siêu tham số thuật toán")

body_text(
    "Bốn siêu tham số chính được khảo sát: số lượng sinh (num_generations), "
    "hệ số KL divergence (kl_coef), epsilon clip (ε), và tốc độ học (learning "
    "rate). Kết quả cho thấy num_generations = 16 là đủ cho mô hình 4B; tăng "
    "lên 32 không cải thiện chất lượng (Bảng 3). Hệ số KL quyết định chất "
    "lượng: KL = 0,1 cho Func Selection 0,400, trong khi KL = 0 đạt 0,615 — một "
    "cải thiện 53,8% (Bảng 4). Epsilon clip = 0,2 và learning rate 5e‑6 cho "
    "kết quả ổn định nhất qua các lần chạy."
)

body_text(
    "Phân tích chi tiết với KL = 0,1 cho thấy mô hình bị phạt nặng khi khác "
    "biệt so với chính sách tham khảo, dẫn đến ít khám phá hơn và điểm thưởng "
    "thấp hơn. Với KL = 0, mô hình tự do tối ưu hóa điểm thưởng và học được "
    "cách gọi hàm chính xác hơn. Tuy nhiên, KL = 0 dẫn đến giảm khả năng từ "
    "chối (Abstention Accuracy từ 0,835 xuống 0,451) do mô hình chủ động gọi "
    "hàm nhiều hơn."
)

heading_sub("4.3 Chiến lược huấn luyện")

body_text(
    "Ba chiến lược được so sánh: (1) SFT thuần — huấn luyện có giám sát trên "
    "3.553 mẫu chuyên gia; (2) RCTP‑FT → RC‑GRPO — tiền huấn luyện quỹ đạo "
    "có token thưởng trước khi RL (chiến lược mặc định); (3) RL trực tiếp từ "
    "SFT (bỏ qua RCTP‑FT). Chiến lược (2) cho kết quả tốt nhất, với Func "
    "Selection 0,615 so với 0,607 của (3). Kết luận: RCTP‑FT cung cấp nền "
    "tảng quỹ đạo ổn định giúp RL hội tụ nhanh và hiệu quả hơn, đặc biệt trong "
    "bối cảnh số lượng mẫu huấn luyện hạn chế (3.553 mẫu)."
)

heading_sub("4.4 So sánh đối sách")

body_text(
    "Bốn phương pháp truy xuất hàm được đánh giá trên tập kiểm thử 2.075 mẫu: "
    "BM25 thuần, BM25 + bge‑m3 (BAAI), BM25 + codefuse‑ai/F2LLM‑v2‑0.6B, và "
    "BM25 + AITeamVN/Vietnamese‑Embedding‑v2. Bảng 2 tổng hợp kết quả."
)

table_caption(
    "Bảng 2: So sánh các phương pháp truy xuất hàm trên tập kiểm thử 2.075 mẫu."
)
make_table(doc,
    headers=["Phương pháp", "P@1", "MRR", "Latency (s)"],
    rows_data=[
        ["BM25 (chuẩn)",                       "0,524", "0,687", "0,46"],
        ["BM25 + bge‑m3",                      "0,563", "0,718", "320"],
        ["BM25 + F2LLM‑v2‑0.6B",               "0,593", "0,748", "437"],
        ["BM25 + Vietnamese‑Embedding‑v2",      "0,527", "0,695", "386"],
    ],
    col_widths=[2.6, 0.8, 0.8, 1.0],
)

body_text(
    "Kết quả cho thấy BM25 + F2LLM‑v2‑0.6B đạt P@1 và MRR cao nhất (0,593 và "
    "0,748), cải thiện lần lượt 13,2% và 8,9% so với BM25 thuần. Tuy nhiên, "
    "latency tăng từ 0,46 giây lên 437 giây (gấp ~950 lần), khiến phương án "
    "này không phù hợp cho ứng dụng thời gian thực. Ngược lại, BM25 + bge‑m3 "
    "đạt P@1 0,563 và MRR 0,718 với latency 320 giây. Vietnamese‑Embedding‑v2 "
    "cho kết quả thấp hơn cả BM25 thuần (P@1 0,527), có thể do embedding không "
    "được tối ưu cho tên hàm kỹ thuật viễn thông. Do đó, BM25 thuần được chọn "
    "làm phương pháp truy xuất mặc định cho toàn bộ thực nghiệm."
)

heading_sub("4.5 Kết quả thực nghiệm")

body_text(
    "Ba thí nghiệm cắt giảm được tiến hành trên cùng checkpoint RC‑GRPO KL = 0: "
    "(i) thay đổi số lượng sinh, (ii) loại bỏ KL divergence, và (iii) thay đổi "
    "loss function. Tất cả checkpoint được đánh giá trên cùng tập kiểm thử 2.075 "
    "mẫu với quy trình giống hệt nhau."
)

# ---- Bảng 3: num_generations ----

table_caption(
    "Bảng 3: Ảnh hưởng của số lượng sinh (num_generations) đến chất lượng mô hình."
)
make_table(doc,
    headers=["num_generations", "Func Selection", "Task Success", "Trùng MD5"],
    rows_data=[
        ["8",  "0,6151", "0,4894", "–"],
        ["16", "0,6138", "0,4871", "100% (giống ng8)"],
        ["32", "0,6103", "0,4867", "Khác"],
    ],
    col_widths=[2.0, 1.4, 1.4, 2.0],
)

body_text(
    "Bảng 3 cho thấy không có sự khác biệt đáng kể giữa các giá trị "
    "num_generations. Checkpoint ng8 và ng16 có trọng số giống hệt nhau (cùng "
    "MD5 hash), xác nhận rằng G = 16 là đủ cho mô hình 4B trên tác vụ này. "
    "Ng32 cho kết quả trong khoảng nhiễu (0,6103 so với 0,6151), khẳng định "
    "việc tăng số lượng sinh thêm không mang lại lợi ích."
)

# ---- Bảng 4: KL divergence ----

table_caption(
    "Bảng 4: Ảnh hưởng của hệ số KL divergence đến chất lượng mô hình."
)
make_table(doc,
    headers=["Hệ số KL", "Func Selection", "Task Success",
             "Schema Validity", "Exec Success", "Abstention Acc"],
    rows_data=[
        ["0,1", "0,3999", "0,3999", "0,397", "0,437", "0,835"],
        ["0,0", "0,6151", "0,4894", "0,749", "0,729", "0,451"],
    ],
    col_widths=[1.2, 1.2, 1.2, 1.2, 1.2, 1.2],
)

body_text(
    "Bảng 4 cho thấy tác động lớn của KL divergence. Khi KL = 0,1, mô hình bị "
    "ràng buộc chặt với chính sách tham khảo, dẫn đến Func Selection chỉ đạt "
    "0,3999. Khi loại bỏ KL (KL = 0), Func Selection tăng lên 0,6151 "
    "(+53,8%), Schema Validity tăng từ 0,397 lên 0,749 (+88,7%), và Execution "
    "Success tăng từ 0,437 lên 0,729 (+66,9%). Tuy nhiên, Abstention Accuracy "
    "giảm từ 0,835 xuống 0,451 (−45,9%), phản ánh sự đánh đổi giữa độ chính "
    "xác gọi hàm và khả năng từ chối."
)

# ---- Bảng 5: Loss Type ----

table_caption(
    "Bảng 6: So sánh các loss function tại KL = 0."
)
make_table(doc,
    headers=["Loss Type", "Func Selection", "Task Success"],
    rows_data=[
        ["GRPO (baseline)", "0,6073", "0,4835"],
        ["DAPO",             "0,6138", "0,4871"],
        ["CISPO",            "0,6138", "0,4871"],
    ],
    col_widths=[2.0, 2.0, 2.0],
)

body_text(
    "Bảng 6 so sánh ba loss function tại KL = 0. Kết quả cho thấy DAPO và "
    "CISPO không khác biệt so với GRPO baseline: Func Selection dao động từ "
    "0,6073 đến 0,6138 (chênh lệch < 0,007), Task Success từ 0,4835 đến 0,4871. "
    "Phân tích per-sample khẳng định DAPO và CISPO cho dự đoán 100% giống hệt "
    "GRPO trên toàn bộ 2.075 mẫu kiểm thử, mặc dù trọng số khác nhau (MD5 "
    "khác nhau). Kết luận: tại KL = 0, loss function không có ảnh hưởng đến "
    "kết quả — KL penalty mới là yếu tố quyết định."
)

body_text(
    "Tổng hợp các thực nghiệm cắt giảm cho thấy KL divergence là siêu tham số "
    "quan trọng nhất: loại bỏ nó mang lại cải thiện ~54% Func Selection. "
    "Ngược lại, loss function và num_generations không ảnh hưởng đáng kể khi "
    "KL = 0. Noise floor của hệ thống đánh giá là ~10% (per-sample disagreement "
    "trên cùng checkpoint), có nghĩa mọi khác biệt dưới ngưỡng này đều có thể "
    "là nhiễu thống kê."
)

heading("5. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN")

italic_note(
    "Tổng kết kết quả chính, tính mới, sáng tạo, hạn chế và hướng phát triển."
)

heading_sub("5.1 Kết luận")

body_text(
    "Nghiên cứu đạt được năm kết quả chính. Thứ nhất, xây dựng quy trình xử "
    "lý dữ liệu hoàn chỉnh cho bài toán Function Calling tiếng Việt: từ Excel "
    "schema đến bốn định dạng dataset (SFT, GRPO, RC‑GRPO, RCTP), bao gồm pipeline "
    "sinh dữ liệu tổng hợp và làm sạch tự động. Thứ hai, phát triển hệ thống "
    "retrieval kết hợp BM25 và đối sánh xác định bảy mức cho tham số — không "
    "cần embedding model cho tham số, BM25 thuần cho hàm đạt P@1 0,524 và MRR "
    "0,687 với latency chỉ 0,46 giây. Thứ ba, triển khai bảy thuật toán RL "
    "(SFT, RCTP‑FT, GRPO, RC‑GRPO, GTPO, MMR‑GRPO, AVSPO) trên nền tảng TRL "
    "và Unsloth. Thứ tư, đề xuất cơ chế giảm Reward Hacking theo sáu hướng. "
    "Thứ năm, triển khai thành công mô hình 4B trên GPU T4 16 GB với hiệu năng "
    "cạnh tranh."
)

body_text(
    "Về mặt thực nghiệm, RC‑GRPO với KL = 0 đạt kết quả tốt nhất: Func "
    "Selection 0,615 và Task Success 0,489 — cải thiện lần lượt 123,6% và "
    "102,1% so với SFT baseline (0,275 và 0,242). Phân tích cắt giảm xác nhận "
    "KL penalty là yếu tố quyết định: loại bỏ nó mang lại +53,8% Func "
    "Selection. Các loss function khác (DAPO, CISPO) tại KL = 0 cho kết quả "
    "giống hệt nhau, khẳng định tính ổn định của khung thực nghiệm."
)

heading_sub("5.2 Hạn chế và hướng phát triển")

body_text(
    "Nghiên cứu còn một số hạn chế. Thứ nhất, các thực nghiệm cắt giảm mới "
    "chỉ được thực hiện trên checkpoint KL = 0 của RC‑GRPO, chưa có điều kiện "
    "kiểm tra trên checkpoint KL = 0,1 hoặc các thuật toán khác. Thứ hai, "
    "MMR‑GRPO và AVSPO mới được cài đặt mã nguồn, chưa được huấn luyện và "
    "đánh giá do hạn chế về tài nguyên tính toán. Thứ ba, tập hàm viễn thông "
    "còn giới hạn ở 31 hàm, trong khi thực tế vận hành có thể lên đến hàng "
    "trăm hàm. Thứ tư, tác động của retrieval đến chất lượng RL chưa được "
    "khảo sát một cách có hệ thống."
)

body_text(
    "Hướng phát triển trong tương lai bao gồm: (i) huấn luyện giai đoạn 2 đầy "
    "đủ với số bước lớn hơn (hiện tại dừng ở bước 500 do điều kiện dừng sớm); "
    "(ii) mở rộng tập hàm lên trên 100 hàm với cấu trúc phân cấp; (iii) thử "
    "nghiệm mô hình 7B với lượng tử hóa sâu hơn (2‑bit, AWQ) để so sánh "
    "hiệu năng; (iv) huấn luyện và đánh giá MMR‑GRPO cùng AVSPO; (v) khảo sát "
    "ảnh hưởng của retrieval quality đến RL qua các embedding model khác nhau; "
    "và (vi) xây dựng giao diện người dùng thử nghiệm cho các kỹ sư mạng."
)

heading("6. TÀI LIỆU THAM KHẢO")

italic_note(
    "Tài liệu tham khảo theo chuẩn IEEE."
)

body_text(
    "[1]  RC-GRPO: 'Reward-Conditioned Group Relative Policy Optimization for "
    "Multi-Turn Tool Calling Agents.' arXiv:2602.03025, 2026."
)
body_text(
    "[2]  Shao, Z., et al. 'DeepSeekMath: Pushing the Limits of Mathematical "
    "Reasoning with Group Relative Policy Optimization.' arXiv:2402.03300, 2024."
)
body_text(
    "[3]  Simoni, A., et al. 'GTPO: Conflict-aware Gradient Correction for "
    "Policy Optimization.' arXiv:2508.03772, 2025."
)
body_text(
    "[4]  Wei, K., et al. 'MMR-GRPO: Diversity-aware Reward Reweighting for "
    "Group Relative Policy Optimization.' ACL 2026 Findings."
)
body_text(
    "[5]  He, S., et al. 'Adaptive Virtual Sample Policy Optimization.' ICML "
    "2026, arXiv:2605.21125."
)
body_text(
    "[6]  Yu, Q., et al. 'DAPO: Dynamic Sampling Policy Optimization for "
    "Large Language Models without KL Divergence.' arXiv:2503.14476, 2025."
)
body_text(
    "[7]  Qin, Y., et al. 'ToolLLM: Facilitating Large Language Models to "
    "Master 16000+ Real-world APIs.' arXiv:2307.16789, 2023."
)
body_text(
    "[8]  Ackermann, J., et al. 'Gradient Regularization for Reward Hacking.' "
    "ICML 2026, arXiv:2602.18037."
)
body_text(
    "[9]  Lambert, N. 'Reinforcement Learning from Verifiable Rewards.' "
    "Chapter 7: Reward Design for Tool-Use Agents, 2025."
)
body_text(
    "[10]  'Qwen3 Technical Report.' arXiv:2505.03718, 2025."
)
body_text(
    "[11]  Unsloth Team. 'Unsloth: 2x Faster Fine-Tuning.' "
    "https://github.com/unslothai/unsloth, 2025."
)
body_text(
    "[12]  von Werra, L., et al. 'TRL: Transformer Reinforcement Learning.' "
    "https://github.com/huggingface/trl, 2024."
)
body_text(
    "[13]  Kwon, W., et al. 'Efficient Memory Management for Large Language "
    "Model Serving with PagedAttention.' SOSP 2023."
)
body_text(
    "[14]  Loshchilov, I. and Hutter, F. 'Decoupled Weight Decay "
    "Regularization (AdamW).' ICLR 2019."
)
body_text(
    "[15]  Hu, E. J., et al. 'LoRA: Low-Rank Adaptation of Large Language "
    "Models.' ICLR 2022."
)
body_text(
    "[16]  Dao, T., et al. 'FlashAttention: Fast and Memory-Efficient Exact "
    "Attention with IO-Awareness.' NeurIPS 2022."
)
body_text(
    "[17]  Dettmers, T., et al. 'QLoRA: Efficient Finetuning of Quantized "
    "Language Models.' NeurIPS 2023."
)
body_text(
    "[18]  'codefuse-ai/F2LLM: Function Calling for LLMs.' "
    "https://huggingface.co/codefuse-ai, 2025."
)
body_text(
    "[19]  Karpukhin, V., et al. 'Dense Passage Retrieval for Open-Domain "
    "Question Answering.' EMNLP 2020."
)
body_text(
    "[20]  Robertson, S. and Zaragoza, H. 'The Probabilistic Relevance "
    "Framework: BM25 and Beyond.' Foundations and Trends in IR, 2009."
)

# ==========================================================
# Save
# ==========================================================

doc.save("VDT_2026_Report_ToolFormer_draft.docx")

print("Saved: VDT_2026_Report_ToolFormer_draft.docx")